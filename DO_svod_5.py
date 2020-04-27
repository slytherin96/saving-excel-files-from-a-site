# -*- coding: utf-8 -*-

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.by import By
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
from selenium.webdriver.ie.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

import datetime
import calendar
import win32com.client
import win32com
import xlrd
import os
#import pandas as pd
import time
import shutil
from win32com.client import Dispatch
import numpy as np
from tkinter import *
import tkinter.messagebox as box
import keyboard
from openpyxl import load_workbook
import docx
from docx.shared import Pt
import docx.section
import pyautogui

#login, password=open('логин и пароль.txt','r').read().split('\n')
load=xlrd.open_workbook(os.getcwd()+'/settings/settings.xlsx').sheet_by_index(0)
login=load.cell_value(0,1)
password=load.cell_value(1,1)
path=load.cell_value(2,1)

def get_week_of_month(year, month, day):#номер недели месяца
    x = np.array(calendar.monthcalendar(year, month))
    week_of_month = np.where(x==day)[0][0]+1
    return(week_of_month)
def num_day(e):# для выбора дня недели, выбирает название элмента из 42 доступных
    lst=[['v1','v2','v3','v4','v5','v6','v7'],['v8','v9','v10','v11','v12','v13','v14'],
      ['v15','v16','v17','v18','v19','v20','v21'],['v22','v23','v24','v25','v26','v27','v28'],
      ['v29','v30','v31','v32','v33','v34','v35'],['v36','v37','v38','v39','v40','v41','v42']]
    startDate = str(e)#первый день месяца
    startDate_datetime = datetime.datetime.strptime( startDate, '%d%m%Y' )
    index=int(startDate_datetime.strftime('%m'))#номер месяца
    year=int(startDate_datetime.strftime('%Y'))#год
    day=int(startDate_datetime.strftime('%d'))#
    dic={1:0,2:1,3:2,4:3,5:4,6:5,0:6}#день недели
    week=int(startDate_datetime.strftime('%w'))
    week=dic[week]
    num_week=get_week_of_month(year,index,day)-1
    return lst[num_week][week]
def dat(e):#вводит нужную дату в поле, написано не целиком а через сложение строк, по другому не берет из за кавычек
    startDate = str(e)#первый день месяца
    startDate_datetime = datetime.datetime.strptime( startDate, '%d%m%Y' )
    a='arguments[0].setAttribute('
    a+="'"
    a+='value'
    a+="'"
    a+=",'"
    b=startDate_datetime.strftime('%d.%m.%Y')
    a+=str(b)
    a+="')"
    return a
def remove_xls():#удаление xls файлов в папке с программой
    file=os.getcwd()
    files = os.listdir(file)
    data=[]
    for i in files:
        if i.find('xl')!=-1:
            if i.find('xlsm')==-1:  
                data.append(i)
    for c in data:
        full_path = os.path.join(file, c)
        if os.path.isfile(full_path):
            os.remove(full_path)
def remove_contents(path):#удаление xl файлов в определенной папке
    file=os.getcwd()
    file=file+'\\сводный\\'+path
    files = os.listdir(file)
    data=[]
    for i in files:
        if i.find('xl')!=-1:
            data.append(i)
    for c in data:
        full_path = os.path.join(file, c)
        if os.path.isfile(full_path):
            os.remove(full_path)
def remove_contents_last(path):#удаление xl файлов в определенной папке
    file=os.getcwd()
    file=file+'\\сводный\\'+path+'\\PrevYear'
    files = os.listdir(file)
    data=[]
    for i in files:
        if i.find('xl')!=-1:
            data.append(i)
    for c in data:
        full_path = os.path.join(file, c)
        if os.path.isfile(full_path):
            os.remove(full_path)
def databeg(e):#выбор нужной даты, начало периодна
    ele =driver.find_element_by_id('idDtbeg')
    driver.execute_script(str(dat(e)), ele)
    driver.find_element_by_id('idDtbeg').click()
    driver.find_element_by_id(num_day(e)).click()
def dataend(e):):#выбор нужной даты, конец периодна
    ele =driver.find_element_by_id('idDtend')
    driver.execute_script(str(dat(e)), ele)
    driver.find_element_by_id('idDtend').click()
    driver.find_element_by_id(num_day(e)).click()
def month_data(e):#выбор первого дня месяца из исходной даты
    startDate = str(e)
    startDate_datetime = datetime.datetime.strptime( startDate, '%d%m%Y' )
    month=str(startDate_datetime.strftime('%m'))#номер месяца
    year=str(startDate_datetime.strftime('%Y'))#год
    return '01'+month+year
def save_excel_and_list_book():# функцция сохранения excel файла через макрос в excel'е
    file=os.getcwd()
    xl=win32com.client.Dispatch("Excel.Application")
    xl.Workbooks.Open(Filename=file+"\\Макрос сохранения.xlsm",ReadOnly=1)
    xl.Application.Run("GetAllXLApp")
#    xl.Workbooks(1).Close(SaveChanges=0)
    xl.Application.Quit()
    
    files = os.listdir(path=".")
    data=[]
    data1=[]
    for i in files:
        if i.find('xl')!=-1:
            data.append(i)
    for i in data:
        if i.find('Книг')!=-1:
            data1.append(str(i))
    data1=sorted(data1,reverse=False)
    return data1
def formatt(my_xlsx_excel_file):#меняет формат сохраненного файла xlsx на xls
    #меняем формат
    xl = Dispatch('Excel.Application')
    wb = xl.Workbooks.Add(my_xlsx_excel_file)
    wb.SaveAs(my_xlsx_excel_file[:-1], FileFormat=56)
    os.remove(my_xlsx_excel_file)
    xl.Quit()
#####################################################	
def sentReply(to_address, subject, do=''):#данная функция создает письмо с вложенным в него word файлом
    #создаем документ
    global now_time
    doc = docx.Document()
    sections = doc.sections
    #вставляем текст в документ
    for section in sections:
        section.left_margin = (400000)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    doc.add_paragraph(now.strftime('%d.%m.%Y')+' в '+now_time+' справка ' + do + ' была сохранена с временем ' + startDate_datetime.strftime('%H:%M') + '.').style = doc.styles['Normal']
    #вставляем избражение
    doc.add_picture(path + do + ' ' + now.strftime('%d.%m.%Y') + '.png', width = docx.shared.Cm(20))
    doc.save(path + do + ' ' + now.strftime('%d.%m.%Y')+'.docx')
    os.remove(path+ do + ' ' + now.strftime('%d.%m.%Y') + '.png')
    # инициализируем объект outlook
    olk = win32com.client.Dispatch("Outlook.Application")
    Msg = olk.CreateItem(0)   
    # формируем письма, выставляя адресата, тему и текст
    Msg.To = to_address
    Msg.GetInspector
    Msg.Subject = subject # добавляем RE в тему
    Msg.Attachments.Add(Source=path + do + ' ' + now.strftime('%d.%m.%Y')+'.docx')
    Msg.display()

##################################################
def save_excel_DO13s(a):#функция сохранения файла и перемещение в нужную папку
    file=os.getcwd()
    data1=save_excel_and_list_book()
    data2=data1
    if len(data2)==0:
        d=2/0
    else:
        os.rename(data1[0], a+".xlsx")
        shutil.move(a+".xlsx", "сводный\\ДО13\\"+a+".xlsx")
        my_xlsx_excel_file=file+"\\сводный\\ДО13\\"+a+".xlsx"
        formatt(my_xlsx_excel_file)  
##################################################
def save_excel_DO13s_last(a):#функция сохранения файла и перемещение в нужную папку
    file=os.getcwd()
    data1=save_excel_and_list_book()
    data2=data1
    if len(data2)==0:
        d=2/0
    else:
        os.rename(data1[0], a+".xlsx")
        shutil.move(a+".xlsx", "сводный\\ДО13\\PrevYear\\"+a+".xlsx")
        my_xlsx_excel_file=file+"\\сводный\\ДО13\\PrevYear\\"+a+".xlsx"
        formatt(my_xlsx_excel_file)  
##################################################
def save_excel_DO11s(a):#функция сохранения файла и перемещение в нужную папку
    file=os.getcwd()
    data1=save_excel_and_list_book()
    data2=data1
    
    if len(data2)==0:
        d=2/0
    else:
        os.rename(data1[0], a+".xlsx")
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/'+ a +'.xlsx').sheet_by_index(0).cell_value(6,0)
        if val=='Не определен':
            wb=load_workbook(a+'.xlsx')
            ws=wb.active
            ws.delete_rows(7,amount=1)
            wb.save(a+'.xlsx')
            os.system("taskkill /f /im  EXCEL.EXE")
        shutil.move(a+".xlsx", "сводный\\ДО11\\"+a+".xlsx")
        my_xlsx_excel_file=file+"\\сводный\\ДО11\\"+a+".xlsx"
        formatt(my_xlsx_excel_file)  
##################################################
def save_excel_DO11s_last(a):#функция сохранения файла и перемещение в нужную папку
    file=os.getcwd()
    data1=save_excel_and_list_book()
    data2=data1
    if len(data2)==0:
        d=2/0
    else:
        os.rename(data1[0], a+".xlsx")
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/'+ a +'.xlsx').sheet_by_index(0).cell_value(6,0)
        if val=='Не определен':
            wb=load_workbook(a+'.xlsx')
            ws=wb.active
            ws.delete_rows(7,amount=1)
            wb.save(a+'.xlsx')
            os.system("taskkill /f /im  EXCEL.EXE")
        shutil.move(a+".xlsx", "сводный\\ДО11\\PrevYear\\"+a+".xlsx")
        my_xlsx_excel_file=file+"\\сводный\\ДО11\\PrevYear\\"+a+".xlsx"
        formatt(my_xlsx_excel_file)  
##########################################################
def save_excel_DO25s(a):#функция сохранения файла и перемещение в нужную папку
    file=os.getcwd()
    data1=save_excel_and_list_book()
    data2=data1
    if len(data2)==0:
        d=2/0
    else:
        os.rename(data1[0], a+".xlsx")
        shutil.move(a+".xlsx", "сводный\\ДО25\\"+a+".xlsx")
        my_xlsx_excel_file=file+"\\сводный\\ДО25\\"+a+".xlsx"
        formatt(my_xlsx_excel_file)   
################################################## 
def save_excel_DO25s_last(a):#функция сохранения файла и перемещение в нужную папку
    file=os.getcwd()
    data1=save_excel_and_list_book()
    data2=data1
    if len(data2)==0:
        d=2/0
    else:
        os.rename(data1[0], a+".xlsx")
        shutil.move(a+".xlsx", "сводный\\ДО25\\PrevYear\\"+a+".xlsx")
        my_xlsx_excel_file=file+"\\сводный\\ДО25\\PrevYear\\"+a+".xlsx"
        formatt(my_xlsx_excel_file)   
##################################################  
def save_excel_pr11s(a):#функция сохранения файла и перемещение в нужную папку
    file=os.getcwd()
    data1=save_excel_and_list_book()
    data2=data1
    if len(data2)==0:
        d=2/0
    else:
        os.rename(data1[0], a+".xlsx")
        shutil.move(a+".xlsx", "сводный\\analitic\\"+a+".xlsx")
        my_xlsx_excel_file=file+"\\сводный\\analitic\\"+a+".xlsx"
        formatt(my_xlsx_excel_file)      
###########################################################   
a=0  
def DO13s():
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_xls()#удаляем файлы из папки с программой
    remove_contents('ДО13')#удаляем файлы в нужной папке
	#функция выбора даты
    def databeg(e):
        ele =driver.find_element_by_id('idDtbeg')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtbeg').click()
        driver.find_element_by_id(num_day(e)).click()
    def dataend(e):
        ele =driver.find_element_by_id('idDtend')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtend').click()
        driver.find_element_by_id(num_day(e)).click()
    global a
    e1=entry1.get()#запрашиваем дату
	#входим на сайт
    opts=Options()
    opts.ignore_protected_mode_settings=True
    opts.ignore_zoom_level=True
    driver = webdriver.Ie('IEDriverServer.exe',ie_options=opts)
    keyboard.press_and_release('ctrl + 0, space')
    time.sleep(0.5)
    keyboard.press_and_release('ctrl + -, space')
    driver.set_window_size(1750, 900)
    driver.get('http...')
    driver.find_elements_by_id('username')[0].send_keys(login)
    driver.find_element_by_id('password').send_keys('')
    driver.find_element_by_id('password').clear()
    driver.find_element_by_id('password').send_keys(password)
    
    driver.find_elements_by_class_name('btn-submit')[0].click()
    #выбираем нужное значение из поля
    a=0
    def zap():
        global a
    #time.sleep(10)
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка')
            a+=1
            time.sleep(0.5)
            zap()
    zap()
    #ждем нужного времени обновления нас сайте. Ожидание сделано через try except и ошибку d=2/0
    a=0
    startDate_datetime=None
    now=None
    tim=None
    now_time=None
    def times():
            global a
            global startDate_datetime
            global now
            global tim
            global now_time
        #time.sleep(10)
            try:
                driver.switch_to.default_content()
                iframe = driver.find_elements_by_tag_name('frame')[0]
                driver.switch_to.frame(iframe)
                startDate=driver.find_element_by_xpath("//label[@id='id_currDT']").text
                startDate_datetime = datetime.datetime.strptime(startDate, '%d.%m.%Y %H:%M' )
                now = datetime.datetime.now()
                now_time=now.strftime('%H:%M')
                today = now.replace(hour=2, minute=30, second=0, microsecond=0)
                if now<today:
                    if startDate_datetime.strftime('%H:%M')!='01:30':
                        d=2/0
                elif startDate_datetime.strftime('%H:%M')=='01:30':
                    pass
                elif startDate_datetime.strftime('%H:%M')!='01:30':
                    #сохранение изображения
                    pyautogui.screenshot(path+'ДО-13 ' + now.strftime('%d.%m.%Y') + '.png')
                
            except:
                
                if a==100:
                    driver.close()
                    sys.exit()
                a+=1
                time.sleep(30)
                driver.get('http...')
                times()
    if var_1.get() == 1: times()
    
    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
#########################Далее во всех функция идет выбор нужных данных и сохранения их в файлы###########################
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s('s_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
   
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s('sd_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения sd_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=2&id=App678')
    os.system("taskkill /f /im  EXCEL.EXE")
    time.sleep(7)
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=2&id=App678')
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s('pr68s_pas_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения pr68s_pas_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=0&id=App9')
    os.system("taskkill /f /im  EXCEL.EXE")
    time.sleep(7)
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=0&id=App9')
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s('pr9s_pas_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения pr9s_pas_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    os.system("taskkill /f /im  EXCEL.EXE")
    e=month_data(e1)
    databeg(e)
   
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=0&id=App9')
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s('pr9m_pas_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения pr9m_pas_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=0&id=Do13Summary')
    time.sleep(7)
    os.system("taskkill /f /im  EXCEL.EXE")
    
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=0&id=Do13Summary')
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s('m_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения m')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    os.system("taskkill /f /im  EXCEL.EXE")
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s('md_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения md_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=2&id=App678')
    os.system("taskkill /f /im  EXCEL.EXE")
    time.sleep(7)
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=2&id=App678')
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s('pr68m_pas_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения pr68m_pas_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    #вовзращаем масштаб браузера
    keyboard.press_and_release('ctrl + 0, space')
    driver.close()
    
    #копируем файлы из локальной папки на сетевой диск(имя сетевой папки и диска заменено на несуществующий)
    def copy_file():
        file=os.getcwd()
        a=file+'\сводный\ДО13'
        b=r'W:\papka'
        os.system("copy /Y " + a + "\m_"+e1+".xls " + b + "\m_"+e1+".xls "
                  " & copy /Y " + a + "\s_"+e1+".xls " + b + "\s_"+e1+".xls "
                  " & copy /Y " + a + "\sd_"+e1+".xls " + b + "\sd_"+e1+".xls "
                 " & copy /Y " + a + "\md_"+e1+".xls " + b + "\md_"+e1+".xls "
                  " & copy /Y " + a + "\sd_"+e1+".xls " + b + "\sd_"+e1+".xls "
                  " & copy /Y " + a + "\pr68s_pas_"+e1+".xls " + b + "\pr68s_pas_"+e1+".xls "
                  " & copy /Y " + a + "\pr68m_pas_"+e1+".xls " + b + "\pr68m_pas_"+e1+".xls "
                  " & copy /Y " + a + "\pr9s_pas_"+e1+".xls " + b + "\pr9s_pas_"+e1+".xls "
                  " & copy /Y " + a + "\pr9m_pas_"+e1+".xls " + b + "\pr9m_pas_"+e1+".xls "
                  )
    if var_2.get() == 1: copy_file()
	#проверка правильности сохранения файлов
    try:
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/сводный/ДО13/sd_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        val1=xlrd.open_workbook(file+'/сводный/ДО13/sd_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-20:]
        if val!='...' or val1!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 sd_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/md_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        val1=xlrd.open_workbook(file+'/сводный/ДО13/md_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-20:]
        if val!='...' or val1!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 md_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/s_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/m_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 m_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/md_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-17:]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 md_'+e1+'.xls') 
        val=xlrd.open_workbook(file+'/сводный/ДО13/pr68s_pas_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-27]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 pr68s_pas_'+e1+'.xls') 
        val=xlrd.open_workbook(file+'/сводный/ДО13/pr68m_pas_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-27]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 pr68m_pas_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/pr9s_pas_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-24]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 pr9s_pas_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/pr9m_pas_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-24]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 pr9m_pas_'+e1+'.xls')
    except:
        box.showerror('Сбор ДО','Проверка файлов не завершена')
    
    os.system("taskkill /f /im  EXCEL.EXE")
    os.system("taskkill /f /im  IEDriverServer.exe")
    box.showinfo("Сохранение ДО", "Сохранение выполнено")
    file=os.getcwd()
    os.startfile(r'W:\papka')
	#если справки были сохранены с другим временем, то создается письмо
    if tim!='01:30':
        sentReply('email', 'ДО-13 не обновилось', do='ДО-13')
########################################################
a=0
def DO11s():
    global path
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_xls()
    remove_contents('ДО11')
    def databeg(e):
        ele =driver.find_element_by_id('idDtbeg')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtbeg').click()
        driver.find_element_by_id(num_day(e)).click()
    def dataend(e):
        ele =driver.find_element_by_id('idDtend')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtend').click()
        driver.find_element_by_id(num_day(e)).click()
    global a
    e1=entry2.get()
    opts=Options()
    opts.ignore_protected_mode_settings=True
    opts.ignore_zoom_level=True
    driver = webdriver.Ie('IEDriverServer.exe',ie_options=opts)
    keyboard.press_and_release('ctrl + 0, space')
    time.sleep(0.5)
    keyboard.press_and_release('ctrl + -, space')
    driver.set_window_size(1750, 900)

    driver.get('http...')
    driver.find_elements_by_id('username')[0].send_keys(login)
    driver.find_element_by_id('password').send_keys('')
    driver.find_element_by_id('password').clear()
    driver.find_element_by_id('password').send_keys(password)
    
    driver.find_elements_by_class_name('btn-submit')[0].click()
  
    
    a=0
    
    def zap():
        global a
    #time.sleep(10)
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка')
            a+=1
            time.sleep(0.5)
            zap()
    zap()
    
    a=0
    startDate_datetime=None
    now=None
    tim=None
    now_time=None
    def times():
            global a
            global startDate_datetime
            global now
            global tim
            global now_time
        #time.sleep(10)
            try:
                driver.switch_to.default_content()
                iframe = driver.find_elements_by_tag_name('frame')[0]
                driver.switch_to.frame(iframe)
                startDate=driver.find_element_by_xpath("//label[@id='id_currDT']").text
                startDate_datetime = datetime.datetime.strptime(startDate, '%d.%m.%Y %H:%M' )
                tim=startDate_datetime.strftime('%H:%M')
                now = datetime.datetime.now()
                now_time=now.strftime('%H:%M')
                today = now.replace(hour=2, minute=30, second=0, microsecond=0)
                if now<today:
                    if startDate_datetime.strftime('%H:%M')!='01:30':
                        d=2/0
                elif startDate_datetime.strftime('%H:%M')=='01:30':
                    pass
                elif startDate_datetime.strftime('%H:%M')!='01:30':
                    #сохранение изображения
                    pyautogui.screenshot(path+'ДО-11 ' + now.strftime('%d.%m.%Y') + '.png')
            except:
                
                if a==100:
                    driver.close()
                    sys.exit()
                a+=1
                time.sleep(30)
                driver.get('http...')
                times()
    if var_1.get() == 1: times()
    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()

    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s('s_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s('sd_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения sd_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=2&id=App678')
    os.system("taskkill /f /im  EXCEL.EXE")
    time.sleep(7)
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=2&id=App678')
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 60).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s('pr14s_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения pr14s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    e=month_data(e1)
    databeg(e)
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=2&id=App678')
    a=0
    def zap():
        global a
        try:
            time.sleep(4)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 60).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s('pr14m_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения pr14m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=0&id=Do11Summary')
    os.system("taskkill /f /im  EXCEL.EXE")
    time.sleep(7)
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=0&id=Do11Summary')
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s('m_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s('md_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения md_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
   
    keyboard.press_and_release('ctrl + 0, space')
    driver.close()
    
    def copy_file():
        file=os.getcwd() 
        c=file+'\сводный\ДО11'
        d=r'W:\papka'
        os.system("copy /Y " + c + "\m_"+e1+".xls " + d + "\m_"+e1+".xls "
                  " & copy /Y " + c + "\s_"+e1+".xls " + d + "\s_"+e1+".xls "
                  " & copy /Y " + c + "\sd_"+e1+".xls " + d + "\sd_"+e1+".xls "
                  " & copy /Y " + c + "\md_"+e1+".xls " + d + "\md_"+e1+".xls "
                  " & copy /Y " + c + "\pr14s_"+e1+".xls " + d + "\pr14s_"+e1+".xls "
                  " & copy /Y " + c + "\pr14m_"+e1+".xls " + d + "\pr14m_"+e1+".xls "
                  )
    if var_2.get() == 1: copy_file()
    
    try:
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/сводный/ДО11/sd_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        val1=xlrd.open_workbook(file+'/сводный/ДО11/sd_'+e1+'.xls').sheet_by_index(0).cell_value(2,0)[-20:]
        if val!='...' or val1!='по Свердловской ж.д.':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения sd_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/md_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        val1=xlrd.open_workbook(file+'/сводный/ДО11/md_'+e1+'.xls').sheet_by_index(0).cell_value(2,0)[-20:]
        if val!='...' or val1!='по Свердловской ж.д.':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения md_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/s_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/m_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения m_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/pr14s_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-27]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr14s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/pr14m_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-27]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr14m_'+e1+'.xls')
    except:
        box.showerror('Сбор ДО','Проверка файлов не завершена')
    
    os.system("taskkill /f /im  IEDriverServer.exe")
    os.system("taskkill /f /im  EXCEL.EXE")
    box.showinfo("Сохранение ДО", "Сохранение выполнено")
    file=os.getcwd()
    os.startfile(r'W:\papka')
    if tim!='01:30':
        sentReply('email', 'ДО-11 не обновилось', do='ДО-11')
##################################################################
a=0

    os.system("taskkill /f /im  EXCEL.EXE")
    remove_xls()
    remove_contents('ДО25')
    def databeg(e):
        ele =driver.find_element_by_id('idDtbeg')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtbeg').click()
        driver.find_element_by_id(num_day(e)).click()
    def dataend(e):
        ele =driver.find_element_by_id('idDtend')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtend').click()
        driver.find_element_by_id(num_day(e)).click()
    global a
    e1=entry3.get()
    opts=Options()
    opts.ignore_protected_mode_settings=True
    opts.ignore_zoom_level=True
    driver = webdriver.Ie('IEDriverServer.exe',ie_options=opts)
    keyboard.press_and_release('ctrl + 0, space')
    time.sleep(0.5)
    keyboard.press_and_release('ctrl + -, space')
    driver.set_window_size(1750, 900)
    driver.get('http...')
    driver.find_elements_by_id('username')[0].send_keys(login)
    driver.find_element_by_id('password').send_keys('')
    driver.find_element_by_id('password').clear()
    driver.find_element_by_id('password').send_keys(password)
    
    driver.find_elements_by_class_name('btn-submit')[0].click()
    time.sleep(3)
 
    a=0
    
    def zap():
        global a
    #time.sleep(10)
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка2')
            a+=1
            time.sleep(0.5)
            zap()
    zap()

    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s('pr1s_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr1s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s('pr3s_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr3s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=11&id=Do25App')
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s('pr2s_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения pr2s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=11&id=Do25App')
    Select(driver.find_element_by_id('id_oper')).select_by_value('1')
    Select(driver.find_element_by_id('id_stan')).select_by_value('STP')
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s('pr4s_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения pr4s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    
    e=month_data(e1)
    databeg(e)
    
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=11&id=Do25App')
    Select(driver.find_element_by_id('id_oper')).select_by_value('1')
    Select(driver.find_element_by_id('id_stan')).select_by_value('STP')
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s('pr4m_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения pr4m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    Select(driver.find_element_by_id('id_oper')).select_by_value('0')
    Select(driver.find_element_by_id('id_stan')).select_by_value('VSE')
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s('pr2m_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr2m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    Select(driver.find_element_by_id('id_rform')).select_by_value('&railroad=076&form=0&id=Do25Summary')
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s('pr1m_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr1m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s('pr3m_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr3m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    keyboard.press_and_release('ctrl + 0, space')
    driver.close()
    def copy_file():
        file=os.getcwd()    
        f=file+'\сводный\ДО25'
        g=r'W:\papka'
        os.system("copy /Y " + f + '\pr1s_'+e1+'.xls ' + g + '\pr1s_'+e1+'.xls '
                  " & copy /Y " + f + '\pr1m_'+e1+'.xls ' + g + '\pr1m_'+e1+'.xls '
                  " & copy /Y " + f + '\pr2s_'+e1+'.xls ' + g + '\pr2s_'+e1+'.xls '
                  " & copy /Y " + f + '\pr2m_'+e1+'.xls ' + g + '\pr2m_'+e1+'.xls '
                  " & copy /Y " + f + '\pr3s_'+e1+'.xls ' + g + '\pr3s_'+e1+'.xls '
                  " & copy /Y " + f + '\pr3m_'+e1+'.xls ' + g + '\pr3m_'+e1+'.xls '
                  " & copy /Y " + f + '\pr4s_'+e1+'.xls ' + g + '\pr4s_'+e1+'.xls '
                  " & copy /Y " + f + '\pr4m_'+e1+'.xls ' + g + '\pr4m_'+e1+'.xls '
                  )
    if var_2.get() == 1: copy_file()
    try:
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/сводный/ДО25/pr1m_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr1m_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/pr1s_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr1s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/pr2s_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[72:-44]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr2s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/pr2m_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[72:-44]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr2m_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/pr3s_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-17:]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr3s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/pr3m_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-17:]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr3m_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/pr4s_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[72:-44]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr4s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/pr4m_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[72:-44]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr4m_'+e1+'.xls')
    except:
        box.showerror('Сбор ДО','Проверка файлов не завершена')
    
    os.system("taskkill /f /im  IEDriverServer.exe")
    os.system("taskkill /f /im  EXCEL.EXE")
    box.showinfo("Сохранение ДО", "Сохранение выполнено")
    file=os.getcwd()
    os.startfile(r'W:\papka')
#######################################################################
a=0
def DO11s_day():
#    os.system("taskkill /f /im  IEDriverServer.exe")
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_xls()
    remove_contents_last('ДО11')
    def databeg(e):
        ele =driver.find_element_by_id('idDtbeg')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtbeg').click()
        driver.find_element_by_id(num_day(e)).click()
    def dataend(e):
        ele =driver.find_element_by_id('idDtend')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtend').click()
        driver.find_element_by_id(num_day(e)).click()
    global a
    e1=entry5.get()
    opts=Options()
    opts.ignore_protected_mode_settings=True
    opts.ignore_zoom_level=True
    driver = webdriver.Ie('IEDriverServer.exe',ie_options=opts)
    keyboard.press_and_release('ctrl + 0, space')
    time.sleep(0.5)
    keyboard.press_and_release('ctrl + -, space')
    driver.set_window_size(1750, 900)

    driver.get('http://romul.gvc.oao.rzd:8080/SASStoredProcess/do?_program=/OCRVFoundation/VGD/Hourly/SP/prig_Frameset&id=Do11Summary&railroad=076')
    driver.find_elements_by_id('username')[0].send_keys(login)
    driver.find_element_by_id('password').send_keys('')
    driver.find_element_by_id('password').clear()
    driver.find_element_by_id('password').send_keys(password)
    
    driver.find_elements_by_class_name('btn-submit')[0].click()
  
    
    a=0
    
    def zap():
        global a
    #time.sleep(10)
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка')
            a+=1
            time.sleep(0.5)
            zap()
    zap()

    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()

    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s_last('s_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s_last('sd_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения sd_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    os.system("taskkill /f /im  EXCEL.EXE")
    e=month_data(e1)
    databeg(e)
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s_last('md_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения md_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    Select(driver.find_element_by_id('id_road')).select_by_value('0')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s_last('m_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
   
    keyboard.press_and_release('ctrl + 0, space')
    driver.close()
    
    def copy_file():
        file=os.getcwd()
        c=file+'\сводный\ДО11\PrevYear'
        d=r'W:\papka\PrevYear'
        os.system("copy /Y " + c + "\m_"+e1+".xls " + d + "\m_"+e1+".xls "
                  " & copy /Y " + c + "\s_"+e1+".xls " + d + "\s_"+e1+".xls "
                  " & copy /Y " + c + "\sd_"+e1+".xls " + d + "\sd_"+e1+".xls "
                  " & copy /Y " + c + "\md_"+e1+".xls " + d + "\md_"+e1+".xls "
                 )
    if var_2.get() == 1: copy_file()
    
    try:
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/sd_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        val1=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/sd_'+e1+'.xls').sheet_by_index(0).cell_value(2,0)[-20:]
        if val!='Отчёт о выполнении расписания движения пригородных' or val1!='по Свердловской ж.д.':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения sd_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/md_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        val1=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/md_'+e1+'.xls').sheet_by_index(0).cell_value(2,0)[-20:]
        if val!='Отчёт о выполнении расписания движения пригородных' or val1!='по Свердловской ж.д.':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения md_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/s_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        if val!='Отчёт о выполнении расписания движения пригородных':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/m_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        if val!='Отчёт о выполнении расписания движения пригородных':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения m_'+e1+'.xls')
    except:
        box.showerror('Сбор ДО','Проверка файлов не завершена')
        
    os.system("taskkill /f /im  IEDriverServer.exe")
    os.system("taskkill /f /im  EXCEL.EXE")
    box.showinfo("Сохранение ДО", "Сохранение выполнено")
    file=os.getcwd()
    os.startfile(r'W:\papka\PrevYear')
    
###################################################################
a=0
def DO25_day():
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_xls()
    remove_contents_last('ДО25')
    def databeg(e):
        ele =driver.find_element_by_id('idDtbeg')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtbeg').click()
        driver.find_element_by_id(num_day(e)).click()
    def dataend(e):
        ele =driver.find_element_by_id('idDtend')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtend').click()
        driver.find_element_by_id(num_day(e)).click()
    global a
    e1=entry4.get()
    opts=Options()
    opts.ignore_protected_mode_settings=True
    opts.ignore_zoom_level=True
    driver = webdriver.Ie('IEDriverServer.exe',ie_options=opts)
    keyboard.press_and_release('ctrl + 0, space')
    time.sleep(0.5)
    keyboard.press_and_release('ctrl + -, space')
    driver.set_window_size(1750, 900)

    driver.get('http...')
    driver.find_elements_by_id('username')[0].send_keys(login)
    driver.find_element_by_id('password').send_keys('')
    driver.find_element_by_id('password').clear()
    driver.find_element_by_id('password').send_keys(password)
    
    driver.find_elements_by_class_name('btn-submit')[0].click()
    time.sleep(3)
   
    a=0
    
    def zap():
        global a
    #time.sleep(10)
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка2')
            a+=1
            time.sleep(0.5)
            zap()
    zap()

    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s_last('pr1s_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr1s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    e=month_data(e1)
    databeg(e)
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s_last('pr1m_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr1m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    keyboard.press_and_release('ctrl + 0, space')
    driver.close()
    
    def copy_file():
        file=os.getcwd()
        f=file+'\сводный\ДО25\PrevYear'
        g=r'W:\papka\PrevYear'
        os.system("copy /Y " + f + '\pr1s_'+e1+'.xls ' + g + '\pr1s_'+e1+'.xls '
                  " & copy /Y " + f + '\pr1m_'+e1+'.xls ' + g + '\pr1m_'+e1+'.xls '
                  )
    if var_2.get() == 1: copy_file()
    try:
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/сводный/ДО25/PrevYear/pr1m_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:]
        if val!='Отчёт о выполнении расписания движения грузовых поездов':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr1m_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/PrevYear/pr1s_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:]
        if val!='Отчёт о выполнении расписания движения грузовых поездов':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr1s_'+e1+'.xls')
    except:
        box.showerror('Сбор ДО','Проверка файлов не завершена')
        
    os.system("taskkill /f /im  IEDriverServer.exe")
    box.showinfo("Сохранение ДО", "Сохранение выполнено")
    file=os.getcwd()
    os.startfile(r'W:\papka\PrevYear')

####################################################################
a=0  
def DO13s_day():
#    os.system("taskkill /f /im  IEDriverServer.exe")
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_xls()
    remove_contents_last('ДО13')
    def databeg(e):
        ele =driver.find_element_by_id('idDtbeg')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtbeg').click()
        driver.find_element_by_id(num_day(e)).click()
    def dataend(e):
        ele =driver.find_element_by_id('idDtend')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtend').click()
        driver.find_element_by_id(num_day(e)).click()
    global a
    e1=entry6.get()
    opts=Options()
    opts.ignore_protected_mode_settings=True
    opts.ignore_zoom_level=True
    driver = webdriver.Ie('IEDriverServer.exe',ie_options=opts)
    keyboard.press_and_release('ctrl + 0, space')
    time.sleep(0.5)
    keyboard.press_and_release('ctrl + -, space')
    driver.set_window_size(1750, 900)
    driver.get('http...')
    driver.find_elements_by_id('username')[0].send_keys(login)
    driver.find_element_by_id('password').send_keys('')
    driver.find_element_by_id('password').clear()
    driver.find_element_by_id('password').send_keys(password)
    
    driver.find_elements_by_class_name('btn-submit')[0].click()
    
    a=0
    def zap():
        global a
    #time.sleep(10)
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка')
            a+=1
            time.sleep(0.5)
            zap()
    zap()
    
    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s_last('s_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
   
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s_last('sd_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения sd_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    os.system("taskkill /f /im  EXCEL.EXE")
    e=month_data(e1)
    databeg(e)
    Select(driver.find_element_by_id('id_road')).select_by_value('76')

    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s_last('md_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения md_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    Select(driver.find_element_by_id('id_road')).select_by_value('0')
    os.system("taskkill /f /im  EXCEL.EXE")
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s_last('m_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()

    keyboard.press_and_release('ctrl + 0, space')
    driver.close()
    
    def copy_file():
        file=os.getcwd() 
        a=file+'\сводный\ДО13\PrevYear'
        b=r'W:\papka\PrevYear'
        os.system("copy /Y " + a + "\m_"+e1+".xls " + b + "\m_"+e1+".xls "
                  " & copy /Y " + a + "\s_"+e1+".xls " + b + "\s_"+e1+".xls "
                  " & copy /Y " + a + "\sd_"+e1+".xls " + b + "\sd_"+e1+".xls "
                  " & copy /Y " + a + "\md_"+e1+".xls " + b + "\md_"+e1+".xls "
                  )
    if var_2.get() == 1: copy_file()
    
    try:
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/sd_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        val1=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/sd_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-20:]
        if val!='...' or val1!='по Свердловской ж.д.':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения sd.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/md_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        val1=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/md_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-20:]
        if val!='...' or val1!='по Свердловской ж.д.':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения md_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/s_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/m_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения m_'+e1+'.xls')
    except:
        box.showerror('Сбор ДО','Проверка файлов не завершена')
    
    os.system("taskkill /f /im  EXCEL.EXE")
    os.system("taskkill /f /im  IEDriverServer.exe")
    box.showinfo("Сохранение ДО", "Сохранение выполнено")
    file=os.getcwd()
    os.startfile(r'W:\papka\PrevYear')
#######################################################################
#########################################################################
def pr11():
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_xls()
    remove_contents('analitic')
    def databeg(e):
        cvartal={1:'01',2:'01',3:'01',4:'04',5:'04',6:'04',7:'07',8:'07',9:'07',10:'10',11:'10',12:'10'}
        #startDate = '12042019'
        startDate_datetime = datetime.datetime.strptime( e, '%d%m%Y' )
        month=str(startDate_datetime.strftime('%m'))
        year=str(startDate_datetime.strftime('%Y'))
        e ='01'+cvartal[int(month)]+year
        ele =driver.find_element_by_id('idDtbeg')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtbeg').click()
        driver.find_element_by_id(num_day(e)).click()
    def dataend(e):
        ele =driver.find_element_by_id('idDtend')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtend').click()
        driver.find_element_by_id(num_day(e)).click()
    global a
    e1=entry1.get()
#    e1='16012020'
    opts=Options()
    opts.ignore_protected_mode_settings=True
    opts.ignore_zoom_level=True
    driver = webdriver.Ie('IEDriverServer.exe',ie_options=opts)
    keyboard.press_and_release('ctrl + 0, space')
    time.sleep(0.5)
    keyboard.press_and_release('ctrl + -, space')
    driver.set_window_size(1750, 900)

    driver.get('http...')
    driver.find_elements_by_id('username')[0].send_keys(login)
    driver.find_element_by_id('password').send_keys('')
    driver.find_element_by_id('password').clear()
    driver.find_element_by_id('password').send_keys(password)

    driver.find_elements_by_class_name('btn-submit')[0].click()
    a=0
    def zap():
        global a
        try:
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            dataend(e1)
            databeg(e1)
    
            
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_pr11s('pr11_pass')

        except:
            if a==60:
                return print('ошибка сохранения pr11_pass')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    driver.get('http...')
    def zap():
        global a
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            
            dataend(e1)
            databeg(e1)
    
            
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_pr11s('pr11_prig')

        except:
            if a==60:
                return print('ошибка сохранения pr11_prig')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    def copy_file():
        file=os.getcwd()
        a=file+'\сводный\\analitic'
        b=r'W:\papka'
        os.system("copy /Y " + a + "\pr11_pass.xls " + b + "\pr11_pass.xls "
                  " & copy /Y " + a + "\pr11_prig.xls " + b + "\pr11_prig.xls "
                  )
    if var_2.get() == 1: copy_file()
    try:
        file=os.getcwd()
        keyboard.press_and_release('ctrl + 0, space')
        val=xlrd.open_workbook(file+'/сводный/analitic/pr11_pass.xls').sheet_by_index(0).cell_value(0,0)[:-161]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr11_pass.xls')
        val=xlrd.open_workbook(file+'/сводный/analitic/pr11_prig.xls').sheet_by_index(0).cell_value(0,0)[:-150]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr11_prig.xls')
    except:
        box.showerror('Сбор ДО','Проверка файлов не завершена')    
    
    driver.close()
    os.system("taskkill /f /im  EXCEL.EXE")
    os.system("taskkill /f /im  IEDriverServer.exe")
    box.showinfo("Сохранение ДО", "Сохранение выполнено")
    file=os.getcwd()
    os.startfile(r'W:\papka')
####################################################################
###################################################################
def DOs_all_day():
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_xls()
    remove_contents_last('ДО11')
    def databeg(e):
        ele =driver.find_element_by_id('idDtbeg')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtbeg').click()
        driver.find_element_by_id(num_day(e)).click()
    def dataend(e):
        ele =driver.find_element_by_id('idDtend')
        driver.execute_script(str(dat(e)), ele)
        driver.find_element_by_id('idDtend').click()
        driver.find_element_by_id(num_day(e)).click()
    global a

    e1=entry5.get()
    opts=Options()
    opts.ignore_protected_mode_settings=True
    opts.ignore_zoom_level=True
    driver = webdriver.Ie('IEDriverServer.exe',ie_options=opts)
    keyboard.press_and_release('ctrl + 0, space')
    time.sleep(0.5)
    keyboard.press_and_release('ctrl + -, space')
    driver.set_window_size(1750, 900)
    #driver = webdriver.Chrome()
    #driver.get(r"C:\Users\User\Desktop\ДО\вход\Менеджер входа в систему SAS.htm")
    #driver.get(r'file:///F:/Менеджер входа в систему SAS.htm')
    driver.get('http...')
    driver.find_elements_by_id('username')[0].send_keys(login)
    driver.find_element_by_id('password').send_keys('')
    driver.find_element_by_id('password').clear()
    driver.find_element_by_id('password').send_keys(password)
    
    driver.find_elements_by_class_name('btn-submit')[0].click()
   
    a=0
    
    def zap():
        global a
    #time.sleep(10)
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка')
            a+=1
            time.sleep(0.5)
            zap()
    zap()
    
    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()

    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s_last('s_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s_last('sd_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения sd_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    os.system("taskkill /f /im  EXCEL.EXE")
    e=month_data(e1)
    databeg(e)
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s_last('md_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения md_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    Select(driver.find_element_by_id('id_road')).select_by_value('0')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO11s_last('m_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    driver.get('http...')
    driver.get('http...')
    time.sleep(5)
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_contents_last('ДО25')
    a=0
    
    def zap():
        global a
    #time.sleep(10)
        try:
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка2')
            a+=1
            time.sleep(0.5)
            zap()
    zap()

    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s_last('pr1s_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr1s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    os.system("taskkill /f /im  EXCEL.EXE")
    e=month_data(e1)
    databeg(e)
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO25s_last('pr1m_'+e1)
    
        except:
            if a==60:
                return print('ошибка сохранения pr1m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    driver.get('http...')
    driver.get('http...')
    time.sleep(7)
    os.system("taskkill /f /im  EXCEL.EXE")
    remove_contents_last('ДО13')
    a=0
    def zap():
        global a
    #time.sleep(10)
        try:
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            Select(driver.find_element_by_id('id_road')).select_by_value('0')
     #       a=0
        except:
            if a==40:
                return print('ошибка')
            a+=1
            time.sleep(0.5)
            zap()
    zap()
    
    a=0
    def zap():
        global a
        try:
            databeg(e1)
    
            dataend(e1)
    
        except:
            if a==60:
                return print('ошибка выбора даты')
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s_last('s_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения s_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
   
    Select(driver.find_element_by_id('id_road')).select_by_value('76')
    os.system("taskkill /f /im  EXCEL.EXE")
    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s_last('sd_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения sd_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    
    os.system("taskkill /f /im  EXCEL.EXE")
    e=month_data(e1)
    databeg(e)
    Select(driver.find_element_by_id('id_road')).select_by_value('76')

    
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s_last('md_'+e1)
        except:
            if a==60:
                return print('ошибка сохранения md_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    Select(driver.find_element_by_id('id_road')).select_by_value('0')
    os.system("taskkill /f /im  EXCEL.EXE")
    a=0
    def zap():
        global a
        try:
            time.sleep(2)
            driver.find_elements_by_xpath("(//label[@id='excelLabel'])")[1].click()
            time.sleep(1)
            driver.switch_to.default_content()
            iframe1 = driver.find_elements_by_tag_name('frame')[1]
            driver.switch_to.frame(iframe1)
            
            WebDriverWait(driver, 40).until_not(EC.visibility_of_element_located((By.XPATH, "//span[@id='progressMessage']")))
            time.sleep(1)
            
            driver.switch_to.default_content()
            iframe = driver.find_elements_by_tag_name('frame')[0]
            driver.switch_to.frame(iframe)
            save_excel_DO13s_last('m_'+e1)

        except:
            if a==60:
                return print('ошибка сохранения m_'+e1)
            a+=1
    #        time.sleep(0.01)
            zap()
    zap()
    keyboard.press_and_release('ctrl + 0, space')
    driver.close()
    
    def copy_file():
        file=os.getcwd()
        a=file+'\сводный\ДО13\PrevYear'
        b=r'W:\papka\PrevYear'
        c=file+'\сводный\ДО11\PrevYear'
        d=r'W:\papka\PrevYear'
        f=file+'\сводный\ДО25\PrevYear'
        g=r'W:\papka\PrevYear'
        os.system("copy /Y " + a + "\m_"+e1+".xls " + b + "\m_"+e1+".xls "
                  " & copy /Y " + a + "\s_"+e1+".xls " + b + "\s_"+e1+".xls "
                  " & copy /Y " + a + "\sd_"+e1+".xls " + b + "\sd_"+e1+".xls "
                  " & copy /Y " + a + "\md_"+e1+".xls " + b + "\md_"+e1+".xls "
                  " & copy /Y " + c + "\m_"+e1+".xls " + d + "\m_"+e1+".xls "
                  " & copy /Y " + c + "\s_"+e1+".xls " + d + "\s_"+e1+".xls "
                  " & copy /Y " + c + "\sd_"+e1+".xls " + d + "\sd_"+e1+".xls "
                  " & copy /Y " + c + "\md_"+e1+".xls " + d + "\md_"+e1+".xls "
                  " & copy /Y " + f + '\pr1s_'+e1+'.xls ' + g + '\pr1s_'+e1+'.xls '
                  " & copy /Y " + f + '\pr1m_'+e1+'.xls ' + g + '\pr1m_'+e1+'.xls '
                  )
    if var_2.get() == 1: copy_file()
    try:
        file=os.getcwd()
        val=xlrd.open_workbook(file+'/сводный/ДО25/PrevYear/pr1m_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr1m_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО25/PrevYear/pr1s_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения pr1s_'+e1+'.xls')
        
            
        val=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/sd_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        val1=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/sd_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-20:]
        if val!='...' or val1!='по Свердловской ж.д.':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 sd_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/md_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        val1=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/md_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[-20:]
        if val!='...' or val1!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 md_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/s_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО13/PrevYear/m_'+e1+'.xls').sheet_by_index(0).cell_value(0,0)[:-51]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО13 m_'+e1+'.xls')
            
        val=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/sd_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        val1=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/sd_'+e1+'.xls').sheet_by_index(0).cell_value(2,0)[-20:]
        if val!='...' or val1!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО11 sd_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/md_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        val1=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/md_'+e1+'.xls').sheet_by_index(0).cell_value(2,0)[-20:]
        if val!='...' or val1!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО11 md_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/s_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО11 s_'+e1+'.xls')
        val=xlrd.open_workbook(file+'/сводный/ДО11/PrevYear/m_'+e1+'.xls').sheet_by_index(0).cell_value(1,0)[:-54]
        if val!='...':
            box.showerror('Сбор ДО','Необходимо проверить правильность сохранения ДО11 m_'+e1+'.xls')
    except:
        box.showerror('Сбор ДО','Проверка файлов не завершена')

    os.system("taskkill /f /im  EXCEL.EXE")
    os.system("taskkill /f /im  IEDriverServer.exe")
    box.showinfo("Сохранение ДО", "Сохранение выполнено")
########################################################################
window = Tk()
window.title( 'Сохранение ДО' )
os.system("taskkill /f /im  IEDriverServer.exe")

entry1 = Entry( window )
entry2 = Entry( window )
entry3 = Entry( window )
entry4 = Entry( window )
entry5 = Entry( window )
entry6 = Entry( window )
entry7 = Entry( window )
window.geometry("910x350")

entry1.insert(0, (datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%d%m%Y"))
entry2.insert(0, (datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%d%m%Y"))
entry3.insert(0, (datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%d%m%Y"))
data25=(datetime.datetime.now()).strftime("%d%m%Y")
data25=data25[:-4]+str(int(data25[-4:])-1)
entry4.insert(0, (data25)) 
entry5.insert(0, (data25))
entry6.insert(0, (data25))
entry7.insert(0, (datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%d%m%Y"))
   

def directory_DO13():
    file=os.getcwd()
    os.startfile(file+'\\сводный\\ДО13')
def directory_DO11():
    file=os.getcwd()
    os.startfile(file+'\\сводный\\ДО11')
def directory_DO25():
    file=os.getcwd()
    os.startfile(file+'\\сводный\\ДО25')
def directory_DO13_last():
    file=os.getcwd()
    os.startfile(file+'\\сводный\\ДО13\\PrevYear')
def directory_DO11_last():
    file=os.getcwd()
    os.startfile(file+'\\сводный\\ДО11\\PrevYear')
def directory_DO25_last():
    file=os.getcwd()
    os.startfile(file+'\\сводный\\ДО25\\PrevYear')
def directory_pr11():
    file=os.getcwd()
    os.startfile(file+'\\сводный\\analitic')
    
label = Label( text = 'ДО-13 Ночь. Введите дату в формате ДДММГГГГ:' )
label1 = Label( text = 'ДО-11 Ночь. Введите дату в формате ДДММГГГГ:' )
label2 = Label( text = 'ДО-25 Ночь. Введите дату в формате ДДММГГГГ:' )
label3 = Label( text = 'ДО-25 День. Введите дату в формате ДДММГГГГ:' )
label4 = Label( text = 'ДО-11 День. Введите дату в формате ДДММГГГГ:' )
label5 = Label( text = 'ДО-13 День. Введите дату в формате ДДММГГГГ:' )
label7 = Label( text = 'pr11. Введите дату в формате ДДММГГГГ:             ' )
label6 = Label( text = 'Во время работы программы будут закрыты процессы Excel' )


btn = Button( window, text = 'Запустить' , command=DO13s)
btn1 = Button( window, text = 'Запустить' , command=DO11s)
btn2 = Button( window, text = 'Открыть папку' , command=directory_DO13)
btn3 = Button( window, text = 'Открыть папку' , command=directory_DO11)
btn4 = Button( window, text = 'Запустить' , command=DO25)
btn5 = Button( window, text = 'Открыть папку' , command=directory_DO25)
btn6 = Button( window, text = 'Запустить' , command=DO25_day)
btn7 = Button( window, text = 'Открыть папку' , command=directory_DO25_last)
btn8 = Button( window, text = 'Запустить' , command=DO11s_day)
btn9 = Button( window, text = 'Открыть папку' , command=directory_DO11_last)
btn10 = Button( window, text = 'Запустить' , command=DO13s_day)
btn11 = Button( window, text = 'Открыть папку' , command=directory_DO13_last)
btn12 = Button( window, text = 'Запустить' , command=pr11)
btn13 = Button( window, text = 'Открыть папку' , command=directory_pr11)
btn14 = Button( window, text = 'Запуск ДО день' , command=DOs_all_day)

btn.grid(row=2,column=3,padx=5,pady=5)
btn1.grid(row=1,column=3,padx=5,pady=5)
btn2.grid(row=2,column=4,padx=5,pady=5)
btn3.grid(row=1,column=4,padx=5,pady=5)
btn4.grid(row=3,column=3,padx=5,pady=5)
btn5.grid(row=3,column=4,padx=5,pady=5)
btn6.grid(row=6,column=3,padx=5,pady=5)
btn7.grid(row=6,column=4,padx=5,pady=5)
btn8.grid(row=4,column=3,padx=5,pady=5)
btn9.grid(row=4,column=4,padx=5,pady=5)
btn10.grid(row=5,column=3,padx=5,pady=5)
btn11.grid(row=5,column=4,padx=5,pady=5)
btn12.grid(row=7,column=3,padx=5,pady=5)
btn13.grid(row=7,column=4,padx=5,pady=5)
btn14.place(x=750,y=140,height=122)

var_1 = IntVar()
var_1.set(1)
book_1 = Checkbutton( window , text = 'Проверка времени\nв ДО-11, 13 Ночь' , variable = var_1 , onvalue = 1, offvalue = 0 )
book_1.place(x=744,y=36)

var_2 = IntVar()
var_2.set(1)
book_2 = Checkbutton( window , text = 'Копирование на\nдиск IHAVGD(W:/)' , variable = var_2 , onvalue = 1, offvalue = 0 )
book_2.place(x=744,y=88)

label.grid(row=2, column=1,padx=5,pady=5)
label1.grid(row=1, column=1,padx=5,pady=5)
label2.grid(row=3, column=1,padx=5,pady=5)
label3.grid(row=6, column=1,padx=5,pady=5)
label4.grid(row=4, column=1,padx=5,pady=5)
label5.grid(row=5, column=1,padx=5,pady=5)
label7.grid(row=7, column=1,padx=5,pady=5)
label6.place(x=5, y=320)


entry1.grid(row=2, column=2,padx=5,pady=5)
entry2.grid(row=1, column=2,padx=5,pady=5)
entry3.grid(row=3, column=2,padx=5,pady=5)
entry4.grid(row=6, column=2,padx=5,pady=5)
entry5.grid(row=4, column=2,padx=5,pady=5)
entry6.grid(row=5, column=2,padx=5,pady=5)
entry7.grid(row=7, column=2,padx=5,pady=5)
window.mainloop()